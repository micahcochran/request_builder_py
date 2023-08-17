# Requires Python 3.5+

# Pseudocode

# Walk through directories.
# Find REQUEST docx/txt files
# Parse REQUEST docx/txt files
# Fill in templates for PC & DRT agenda, publication and letters, and such from request text.

# NOTE: this has nothing to do with the requests python library that makes it a breeze to make HTTP requests.

# TODO: Need to check through the directories and make sure that each one has a REQUEST. Reduces problems.
#    - _find_request_files()
#    - use glob() to create a list of subfolders in the DRT folder
#    - take the results of _find_request_files(), compare the paths to the list of DRT subfolder.  (Portions of the set that are different?)
#
# WISHLIST:
#  - The Application portion should be an interactive editor to generate the request text.
#      - Spell checker
#  - It would be nice to have a each request appear next to the application PDF file for that request.  
#    Or perhaps a link to the PDF file to be able to open the application.
#  - restructure as MVC applicaiton with database
#  - add a GUI mode,  have a Command line argument "--gui" that starts a Graphical User Interface 

####  INSTALLATION  ######################################################
#
# Requires Python 3.6+
# 
# After installed do the following:
# * loguru
#    C:\...> pip install loguru
# * python-docx
#    C:\...> pip install python-docx
# * docxtpl
#    C:\...> pip install docxtpl
#

###  UPDATES  ############################################################
# 
# 2021-02-25  - added folder_exists_or_create(), which asks user to create a folder if the folder doesn't exist
#             - looks into folders to see if all of the cases have REQUEST files
#
# 2021-06-03    - Fixed where it was not printing to the screen.
# version 0.2.1 - Updated legal paper notice to a version that is no longer for COVID.  
#               - Added new lines between items for newpaper publication for better results.
#               - Asks before overwriting PC Notice file

# internal libraries
import datetime as dt
from enum import Enum
from itertools import chain
import os
from pathlib import Path
import sys

if sys.version_info < (3,9):
    from typing import Generator
else:
    # This is the syntax for Python 3.9+
    from collections.abc import Generator

# external libaries
from docx import Document
from docx.shared import Cm
from docxtpl import DocxTemplate, RichText
from loguru import logger


__version__ = '0.2.1'

class UnspecifiedInputError(Exception):
    pass

DEBUG = False

class Request:
    def __init__(self, text, tags={}):
        self.text = text
        self.tags = tags

class Requests:
    """This class locates requests and stores them in memory."""
    def __init__(self, folder: Path):
        self.folder = folder
        self._requestfiles = self._find_request_files(folder)
        # sort _requestfiles based on the folder's number before the period
        # 2. Micah Subdivision
        # 11. Big Annexation
        # The below code  will cause these to sort properly, otherwise 11. would be sorted
        # before 2 because of their ASCII values.
        #
        # For the sorting key, The lambda function isolates the folder's name,
        # splits the portion before the period, and converts that into a numeric value.
        # That numeric values is used to sort the files.
        self._requestfiles.sort(key=lambda x: int(x.parts[-2].split('.')[0]))

        # stores the requests as a list of Request Objects
        self.requests = []
        # list of path object, for folders that do not have request objects
        # TODO: a method is probably a better way to get to this information
        self.folders_without_requests = self._find_folders_without_requests(folder)

# WORKING OLD CODE
#        for fn in self._requestfiles:
#            self.requests.append(self._get_request_text(fn))

        for fn in self._requestfiles:
            req_text, tags = self._get_request_text_with_tags(fn)
            req_obj = Request(req_text, tags)
            self.requests.append(req_obj)

#        if self.requests == []:
#            raise ValueError("No requests files found.")

    def _find_request_files(self, folder: Path):
        """Finds files .docx Word files and .txt files named 'REQUEST'."""
        if os.name == 'nt':
           # this worked fine under Windows because file extensions are case insensitive, 
            docx_files = [fn for fn in folder.rglob('*.docx') if 'REQUEST' in fn.name]
            txt_files = [fn for fn in folder.rglob('*.txt') if 'REQUEST' in fn.name]
        else:
            # Fix for Linux case sensitive filenames -- MDC 2023-08-08
            docx_files = [fn for fn in chain(folder.rglob('*.docx'), folder.rglob('*.DOCX')) if 'REQUEST' in fn.name.upper()]
            txt_files = [fn for fn in chain(folder.rglob('*.txt'), folder.rglob('*.TXT')) if 'REQUEST' in fn.name.upper()]
        return docx_files + txt_files

    def _find_folders_without_requests(self, folder: Path) -> set:
        # get a list of paths for the reqeust files
        folders_with_requests_files = set([reqfile.parent for reqfile in self._requestfiles])

        # create a set of subfolders in this folder
        sub_folders = set([fo for fo in folder.iterdir() if fo.is_dir()])

        return sub_folders - folders_with_requests_files



    # NOTE: UNUSED
    def _get_request_text(self, filename: str) -> str:
        # Note: I think function this might be a little fragile for parsing text.
        # TODO: Could extend development name.
        doc = Document(filename)
        return ''.join([p.text for p in doc.paragraphs])
            

# TODO: use pathlib not strings, add text file support
    def _get_request_text_with_tags(self, filename: str):
        # Note: I think function this might be a little fragile for parsing text.
        # TODO: Could extend development name.
        if filename.suffix.lower() == '.docx':
            doc = Document(filename)
            return self._parse_request_with_variables(doc.paragraphs)
        else:  # assume it is a .txt file
            txtfile = Path(filename)
            with txtfile.open() as fh:
                return self._parse_request_with_variables_txtfile(fh)
                

    # NOTE: THIS CODE IS SPECIFIC TO WORD & DOCX, should be made more general to work with TEXT files
    def _parse_request_with_variables_docxfile(self, paragraphs):
        variable_mode = False
        variables = {}
        request_text = ''
        
        for para in paragraphs:
            if para.text.startswith('---'):
                # the --- sequence toggles variable_mode on/off
                if variable_mode:
                    variable_mode = False
                else:
                    variable_mode = True
                continue

            if variable_mode is True:
                # variable_text_list.append(para.text)
                name, value = para.text.split(':')
                # TODO Could use some robustness for handling empty lines, or throwing out text that doesn't have a variable pair
                # always lower case the variables name
                variables[name.strip().lower()] = value.strip()
            else:
                request_text += para.text
        
        return request_text, variables

    # NOTE: THIS WILL BE VERY SIMILAR TO THE DOCX version above
    # I should probably make a function that I only have to replace the different objects.
    def _parse_request_with_variables_txtfile(self, fh):
        variable_mode = False
        variables = {}
        request_text = ''

        for line in fh:
            if line.startswith('---'):

                # the --- sequence toggles variable_mode on/off
                if variable_mode:
                    variable_mode = False
                else:
                    variable_mode = True
                continue
 
            if variable_mode is True:
                # variable_text_list.append(para.text)
                name, value = line.split(':')
                # TODO Could use some robustness for handling empty lines, or throwing out text that doesn't have a variable pair
                # always lower case the variables name
                variables[name.strip().lower()] = value.strip()
            else:
                request_text += line 

        return request_text, variables

    def item_requires_public_hearing(self, text: str) -> bool:
        # not this was copied from requires_public_hearing.
        # perhaps the classification needs to be property of this class?
        # However a dict[list] really works for the current organization of the agendas.
        keywords = ['preliminary', 'certificate to subdivide', \
                    'replat', 're-plat', 'conditional use', 'annex', 'rezone', 'rezoning']
        return any(kw in text.lower() for kw in keywords)

# TODO: might not be a bad idea to have an override keywords that will otherwise get filter out
# NO_PUBLIC_HEARING YES_PUBLIC_HEARING
# or a header that --- Public Hearing: yes
    def requires_public_hearing(self, idx: int) -> bool:
#        print(f"idx: {idx}")
        if not (0 <= idx <= len(self.requests)):
            raise ValueError("idx value is out of bounds")
        
        # TODO include rezonings and annexations which zone property
        # "any" of these keywords are in the text, meaning the items require a public hearing
        keywords = ['preliminary', 'certificate to subdivide', \
                    'replat', 're-plat', 'conditional use', 'annex', 'rezone', 'rezoning']

#       when this was a text based object
#        return any(kw in self.requests[idx].lower() for kw in keywords)
        return any(kw in self.requests[idx].text.lower() for kw in keywords)

    def requires_city_mailed_notice(self, idx: int) -> bool:
        """
        Any items that requires the City to mail a notice for a public hearing.
        """
#        print(f"idx: {idx}")
        if not (0 <= idx <= len(self.requests)):
            raise ValueError("idx value is out of bounds")
        
        keywords = ['preliminary', 'certificate to subdivide', \
                    'replat', 're-plat']

        # Check if any of the keywords are in the text.
        return any(kw in self.requests[idx].text.lower() for kw in keywords)


    # This is a generator that gives the items that require a public hearing.
    def items_requiring_public_hearing(self) -> Generator:
        for i in range(len(self.requests)):
           if self.requires_public_hearing(i):
               yield self.requests[i]


    # This is a generator that gives the items that require a city mailed notice.
    def items_requiring_city_mailed_notice(self) -> Generator:
        for i in range(len(self.requests)):
           if self.requires_city_mailed_notice(i):
               yield self.requests[i]
    
#    def iterate(self):
#        for i in range(len(self.requests)):
#            yield self.requests[i]

    def classify_cases(self) -> dict:
        """This classifies case for the Planning Commission Agenda."""
        classified = {}
        classified['unclassified'] = []

        # class Case(Enum):
        #     OLDBIZ = '###123'
        #     CERTS = "certificate"
        #     MINOR_SUB = "minor subdivision"
        #     SUB = "major subdivision"
        #     REZONE = "rezone"
        #     ANNEX = "annex"
        #     LOC_CHAR_EXT = "###123"
        #     RULE_AMENDMENT = "###123"
        
        # # put all the keywords into one list
        # keywords = [c.value for c in Case]

        keywords = ['annex', 'rezone', 'certificate', 'minor subdivision',
                    'major subdivision', 'rezone', 'annex']
        for req in self.requests:
            # logger.debug(f'REQ text: {req.text}')
            if any(kw in req.text for kw in keywords):
                for kw in keywords:
                    if kw in req.text:
                        # set up an empty list, if not there
                        classified.setdefault(kw, [])
                        classified[kw].append(req.text)
                        break
            else:
                classified['unclassified'].append(req.text)

        return classified
# kw in self.requests[idx].lower() for kw in keywords

# Does this item need a public hearing?
# Searches if certain key phrases are in that require a public hearing.
# This will occationally get tripped up.
# TODO: might not be a bad idea to have an override keywords that will otherwise get filter out
# NO_PUBLIC_HEARING YES_PUBLIC_HEARING
# def item_requires_public_hearing(text: str) -> bool:
#    requires_pub_hearing_keywords = ['preliminary', 'certificate to subdivide', 'replat', \ 
#                                     're-plat', 'conditional use']
    
#    return requires_pub_hearing_keywords in lower(text) 


# This class generates using the docx templates.
class GenerateTemplates:
    """Generates agendas, letters, and notices based on docx templates."""
# requests is of type list[str]
    def __init__(self, meeting_dates, drt_folder, pc_folder):
        self.templates = 'templates'
        self.meeting_dates = meeting_dates
        self.drt_folder = drt_folder
        self.pc_folder = pc_folder

    def generate_public_hear_form_for_newspaper_legal(self, requests: Requests):
        pub_hearing_requests = []
        # if this becomes more compliated it should be done by a Requests class
        # for req in requests:
        #    if item_requires_public_hearing(req):
        #        pub_hearing_requests.append(req)
        
        # These are Request objects
        ph_list = list(requests.items_requiring_public_hearing())
        # convert Request objects to a list of strings with newlines
        ph_list = [req_obj.text + '\n' for req_obj in ph_list]
        
        context = {
            'pc_meeting_date_str': spelled_out_date(self.meeting_dates.pc),
            'paper_notice_date': self.meeting_dates.paper_notice.isoformat(),
         #   'public_hearing_list': "  This is just for a test " 
#            'public_hearing_list' : RichText('\a'.join(ph_list))
            'public_hearing_list': RichText(self.numbered_list(ph_list))
        }
        doc = DocxTemplate(self.templates + "/PC Notice Template.docx")
        doc.render(context)
        publish_date_str = self.meeting_dates.paper_notice.isoformat() 
        # doc.save(drt_folder + f"\public notice\PC Notice {publish_date_str}.docx")
#        print(f'Public Notice Folder "{pn_folder}" exists: {pn_folder.exists()}')

        pn_folder = self.drt_folder / 'public notice'
        # if the 'public notice' folder doesn't exist, create it
        if pn_folder.is_dir() is False:
            pn_folder.mkdir()
            print(f"Created folder: {pn_folder}")

        notice_file = self.drt_folder / 'public notice' / f'PC Notice {publish_date_str}.docx'

        if file_does_not_exist_or_user_allows_overwriting(notice_file) is True:
            doc.save(notice_file)
            print(f"Wrote file: {notice_file}")


    # Pseudocode
    # 1. classify the items  (NOTE: this level might be overkill.)
    # 2. arrange the agenda using the clasifications
    #     One appraoch would be to have a section for the items

    # TODO THIS DOESN'T WORK.  I RAN OUT OF TIME FOR THIS in November.
    def generate_agenda(self, requests):
        # lighter elements should go first
        weight = {  "old_business": 0,
                    "certificates": 10,
                    "minor subdivision": 20,
                    "subdivsion":30,
                    "rezoning":40,
                    "annexation":45,
                    "location/character/extent":70,
                    "ord/reg amendment":100
        }

        context = {
            'pc_meeting_date_str': spelled_out_date(self.meeting_dates.pc),
            'return_revised_plans_date_str': None # TODO
        }
        doc = DocxTemplate(self.templates + "/PC Agenda Template.docx")
        doc.render(context)


        classified = requests.classify_cases()
        i = 2

        # write the items 
        for key, cases in classified.items():
            # add headings like "REZONING"
            p = doc.add_paragraph('')
            run = p.add_run("\n" + key.upper())
            run.bold = True
            run.underline = True
            for case in cases:
                p = doc.add_paragraph('\n')
                need_ph = requests.item_requires_public_hearing(case)
                if need_ph is True:
                    # Request needs a Public Hearing
                    p.add_run(f'{i}) ')
                    p.paragraph_format.left_indent = Cm(0.5)
                    p = doc.add_paragraph('a) ')
                    run = p.add_run('Public Hearing')
                    run.underline = True
                    run.bold = True
                    p.add_run('. ')
                    p.add_run(case)
                    p.paragraph_format.left_indent = Cm(1)
                    p = doc.add_paragraph('\nb) ')
                    run = p.add_run('Resolution')
                    run.underline = True
                    run.bold = True
                    p.add_run('. ')
                    p.paragraph_format.left_indent = Cm(1)
                else:
                    # Request doe NOT need a Public Hearing
                    p.add_run(f'{i}) ')
                    run = p.add_run('Resolution')
                    run.underline = True
                    run.bold = True
                    p.add_run('. ')
                    p.add_run(case)
                    p.paragraph_format.left_indent = Cm(0.5)
                i += 1
        filename = self.pc_folder / f'GENERATED - PC Agenda - {self.meeting_dates.pc.isoformat()}.docx'
        
        if file_does_not_exist_or_user_allows_overwriting(filename) is True:
            doc.save(filename)
            print(f"Wrote file: {filename}")

    def generate_drt_agenda (self, requests: Requests):
        """Generate Departmental Review Team (DRT) Agenda

        outputs file named:  GENERATED  MONTH YYYY DRT Agenda YYYY-MM-DD.docx"""
        logger.debug(requests.requests)

        # this adds the dates onto the agenda
        context = {
            'drt_date': spelled_out_date_w_weekday(self.meeting_dates.drt),
            'return_revised_plans_date': spelled_out_date_w_weekday(self.meeting_dates.friday_resubmittal)
        }
        doc = DocxTemplate(self.templates + "/DRT Agenda Template.docx")
        doc.render(context)

        # add agemda items 
        # This is a tuple representing the departments which comment in the Departmental Review Team (DRT)
        departments_tuple = ('FIRE', 'WATER', 'ES&CD', 'ELECTRIC', 'GAS', 'CITY ENGINEER', 'MISC', '')
        departments_comment_text = ':\n\n'.join(departments_tuple)
        for n, req in enumerate(requests.requests, start=1):
            p = doc.add_paragraph('')
            p.add_run(f'{n}. {req.text}').bold = True
            doc.add_paragraph(departments_comment_text)
        agenda_fn = drt_folder / f"GENERATED - {self.meeting_dates.drt.strftime('%B %Y')} DRT Agenda - {self.meeting_dates.drt.isoformat()}.docx"
        if file_does_not_exist_or_user_allows_overwriting(agenda_fn) is True:
            # try:
                doc.save(agenda_fn)
            # except PermissionError:
            #    print("PermissionError: Access")
            # TODO needs to catch PermissionError, this happens when you can't open the file.  In case whne you have the other file open in Word or another program.

    def generate_city_mailed_notice(self, requests:Requests):
        mailed_notice_folder = self.drt_folder / 'mailed notice'
        # if the 'mailed notice' folder doesn't exist, create it
        if mailed_notice_folder.is_dir() is False:
            mailed_notice_folder.mkdir()
            print(f"Created folder: {mailed_notice_folder}")

        # commented out COVID related meeting procedures
        # Ask about which template to use for the mailed notice
#        print("Which type of template do you want to use?\n")
        
#        print("R) Regular meeting")
#        print("I) COVID in person meeting")
#        print("V) COVID virtual meeting")
        
#        meeting_selection = ''
#        while meeting_selection not in ['r', 'i', 'v']:
#            meeting_selection = input("?").lower()

        i = 0
        for request_obj in requests.items_requiring_city_mailed_notice():
            # Note: Might be easier to use a json dump from the application PDF files
            
            dev_name = 'Untitled Development'
            # get development name from tags
            if 'short_title' in request_obj.tags: 
                dev_name = request_obj.tags['short_title']

            mailing_date = self.meeting_dates.mailed_notice.isoformat()

            context = {
                'mailing_date': mailing_date,
                'development_name': dev_name,
                'pc_meeting_date': spelled_out_date_w_weekday(self.meeting_dates.pc),  # REDO: Failed in November, fixed in December
                'request_text': request_obj.text 
            }
            if DEBUG is True:
                print(f"DEBUG: context={context}")

            doc = DocxTemplate(self.templates + "/PC mailed notice Template.docx")
            # commented out COVID related meeting proceedures
#            if meeting_selection == 'r':  # Regular Meeting
#                doc = DocxTemplate(self.templates + "\PC mailed notice Template.docx")
#            elif meeting_selection == 'i':  # COVID in person meeting
#                doc = DocxTemplate(self.templates + "\PC mailed notice Template - COVID Inperson.docx")
#            elif meeting_selection == 'v':  # COVID virtual meeting 
#                doc = DocxTemplate(self.templates + "\PC mailed notice Template - COVID virtual.docx")
#            else:
#                raise UnspecifiedInputError("The input function is giving results that it shouldn't be, namely '{meeting_selection}'.  You shouldn't be seeing this.")
            doc.render(context)
            
            # Saves the notice as an number and a development name,
            # in case there is more than one of the same name.
            filename = mailed_notice_folder / f'PC mailed notice {i} - {dev_name} - mail {mailing_date}.docx'
            if file_does_not_exist_or_user_allows_overwriting(filename):
                doc.save(filename)
            i += 1
        print(f'Wrote {i} files to folder: {mailed_notice_folder}')

    def numbered_list(self, req_list):
        s = ''
        for num, item in enumerate(req_list, start=1):
            s += f'{num}. {item}\n'
        return s
        
#    def generate_subdivision_letters(meeting_dates, requests):
#        for req in requests:
#            if item_requires_public_hearing(req):
#                pub_hearing_requests.append(req)

# 
# Class generates all the associated dates based on the year an month.
# These may need to be shifted because of holidays.
# A user interface could override dates.
class MeetingDates:
    """
    Generate all of the associated dates
    """
    def __init__(self, year, month):
        self.year = year
        self.month = month
        # Calculated all of the dates
        self.pc                 = self._calculate_third_day_month(year, month, 1)
        self.friday_resubmittal = self.pc + dt.timedelta(days=-4)  # a Friday
        self.mailed_notice      = self.pc + dt.timedelta(days=-5)
        self.paper_notice       = self.pc + dt.timedelta(days=-10)
        self.drt                = self.pc + dt.timedelta(days=-13)
        self.submittal_deadline = self.pc + dt.timedelta(days=-21) 


    # What is the date of the Third day of the Month?
    # weekday = 1 is Tueday
    # weekday = 3 is Thursday
    def _calculate_third_day_month(self, year: int, month: int, weekday: int) -> dt.date:
        """Calculates the third of a given weekday of the month."""
        # 15th is the lowest possible third day of the month
        third = dt.date(year, month, 15)
        # What day of the week is the 15th?
        w = third.weekday()
        if w != weekday:
            return third.replace(day=15+(weekday - w) % 7)
        return third

    def __str__(self) -> str:
        """Print out all of the Planning Commission related dates"""
        s =   '\n'
        s += f'Planning Commission dates for {self.year}-{self.month:02d}\n'
        s +=  '-----------------------------------------\n'
        s += f'Submittal Date                 {self.submittal_deadline.isoformat()}\n'
        s += f'Departmental Review Team       {self.drt.isoformat()}\n'
        s += f'Newspaper Public Hearings Date {self.paper_notice.isoformat()}\n'
        s += f'Mailed Notice for Subdivisions {self.mailed_notice.isoformat()}\n'
        s += f'Resubmittal Date               {self.friday_resubmittal.isoformat()}\n'
        s += f'Planning Commission Meeting    {self.pc.isoformat()}\n'
        
        return s

# This would return a string for datetime.date(2020, 11, 23) as
# Monday, November 23, 2020
# def weekday_month_day_year(dateobj):
#    return dateobj.strftime('%A, %B %d, %Y')

def spelled_out_date_w_weekday(dateobj: dt.date) -> str:
    """
    >>> spelled_out_date_w_weekday(datetime.date(2020, 11, 23))
    "Monday, November 23, 2020"
    """
    return dateobj.strftime('%A, %B %d, %Y')

def spelled_out_date(dateobj: dt.date) -> str:
    """
    >>> spelled_out_date(datetime.date(2020, 11, 23))
    "November 23, 2020"
    """
    return dateobj.strftime('%B %d, %Y')

# TODO rework to overrides the Document.save method.
# TODO add all feature?
# returns True if file doesn't exist
#              or if the user answers yes
def file_does_not_exist_or_user_allows_overwriting(file_obj):
    """

    returns True if file does not exist, or the user allows overwriting the file
    """
    if file_obj.exists() is False:
        return True
    
    yn = input(f'The file "{file_obj}"" already exists, do you wish to overwrite?  (y)es/[N]o ')
    if yn.lower() == 'y':
        return True
    
    return False

def folder_exists_or_create(path_obj):
    """
    Checks for the existance of a folder.  When it doesn't exist, it prompts user to create the folder.

    returns True if folder exists (or was created during function)
    returns False if folder does not exist
    """  
    if path_obj.is_dir() is True:
        return True

    yn = input(f'The folder "{path_obj}"" does not exist, do wish to create the folder?  (y)es/[N]o ')
    if yn.lower() == 'y':
        path_obj.mkdir()
        print(f"Created folder: {path_obj}")
        return True

    return False

# def gui():
#     """A simplistic Graphical User Interface to allow user to use a dropdown instead of the command line"""

#     raise NotImplementedError("There is still a lot of work to be done on this simple GUI.")
    
#     try:
#         import PySimpleGUI as sg
#     except:
#         print("You will need to install PySimpleGUI.  To install:\n    C:\...> pip install PySimpleGUI") 

#     layout = [[sg.Text("Year Month(##) Report")],
#               [sg.Input(key='year'), sg.Input(key='month'), sg.DropDown(['DATES', 'DRT', 'PC', 'PCNEWS', 'PCMAIL'], key='report')], 
#               [sg.Button('Ok'), sg.Button('Exit')],
#               [sg.Output(size=(50,10), key='-OUTPUT-')] ]
#     while True:
#         (event, value) = window.read()
#         if event == 'Exit' or event == sg.WIN_CLOSED:
#             sys.exit(0)
        
#         if event == 'Ok':
#             year = int(value['year'])
#             month = int(value['month'])
        
#             report = value['report']
#             if report == 'DRT':
#                 pass
    

def usage():
    """Print the help for the program's Command Line Interface"""

    help = """
Usage:
    C:\...> python3 request_builder2.py [date] [Report]

    [date]
        is in YYYY-MM format, so December 2020 is 2020-12
    Report is: 
        DATES - Prints out the meeting dates for a given month
        DRT  - Generate DRT Agenda
        PC  -  Generate Planning Commision Agenda
        PCNEWS - PC Public Hearings newspaper publication
        PCMAIL - Planning Commission Mailings
        ZNGANX - Set Public Hearings for Rezoning/Zoning/Annexation (Does NOT yet work.)
"""
    print(help)

################################################################################################
if __name__ == '__main__':
# This works pretty well for testing.
    if len(sys.argv) == 1:   # No Arguments
        usage()
        sys.exit(10)
    if len(sys.argv) > 3:
        print("Too many arguments.")
        usage()
        sys.exit(10)
    elif len(sys.argv) < 3:
        print("Too few arguments.")
        usage()
        sys.exit(10)

    # parses meeting's year and month
    meeting_ym = dt.datetime.strptime(sys.argv[1], "%Y-%m")

    # Detect possible typo in year argument.
    if meeting_ym.year != dt.date.today().year:
        yn = input('The inputted year is different from this year.  Continue? [y/N] ')
        if yn.lower() != 'y':
            print('Exiting')
            sys.exit(0)

    meeting_dates = MeetingDates(year=meeting_ym.year, month=meeting_ym.month)
    
    report = sys.argv[2]

    # this report is just date calculations that do not rely upon the existance of folders.
    if report == 'DATES':
        print(meeting_dates)
        sys.exit(0)

    drt_folder = Path(f'../DRT/{meeting_dates.drt.isoformat()} DRT')
    if folder_exists_or_create(drt_folder) is False:
        print(f'Path "{drt_folder}" does not exist. Exiting.')
        sys.exit(1)
    requests = Requests(drt_folder)

# DEBUG information
#    logger.debug('FOLDERS WITHOUT REQUESTS')
#    logger.debug(requests.folders_without_requests)

    # Are the any folders without requests files?
    if requests.folders_without_requests != set():
        for folder in requests.folders_without_requests:
            print(folder)
        yn = input("The above folders do not have request files.  Do you wish to continue? [y/N] ")
        if yn.lower() != 'y':
            print('Exiting.')
            sys.exit(1)
    
    # Create the PC main folder for the Planning Commission
    pc_main_folder = Path(f'../PC')
    if folder_exists_or_create(pc_main_folder) is False:
        print(f'Path "{pc_main_folder}" does not exist. Exiting.')
        sys.exit(1)

    pc_folder = Path(f'../PC/{meeting_dates.pc.isoformat()} PC')
    if folder_exists_or_create(pc_folder) is False:
        print(f'Path "{pc_folder}" does not exist. Exiting.')
        sys.exit(1)

    gen_templates = GenerateTemplates(meeting_dates, drt_folder, pc_folder)

    # DRT and PCMAIL works pretty well.
    if report == 'DRT':
         gen_templates.generate_drt_agenda(requests)
    elif report == 'PC':
        gen_templates.generate_agenda(requests)
        print("Note:  Please check the agenda items and reorder items as nessary.")
    elif report == 'PCMAIL':
        # Generates mailed notices
        gen_templates.generate_city_mailed_notice(requests)
    elif report == 'PCNEWS':
        gen_templates.generate_public_hear_form_for_newspaper_legal(requests)
    elif report == 'ZNGANX':
        # assumetions that the legal description in a file named LEGALDES .docx or .txt
        raise NotImplementedError("ZNGANX is not implemented")
    else:
        print(f"The report '{report}' is not a type of report that this software can generate.")
